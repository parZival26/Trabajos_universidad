import docx
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


from chatgpt import chat_gpt

from datetime import date

today = date.today()



def apa_doc(parrafos, nombre, titulo):
    if not isinstance(parrafos, list):
        raise TypeError('El parámetro "parrafos" debe ser una lista')
        

    # Crear documento
    document = docx.Document()

    # Metadatos del documento
    document.add_heading(titulo, 0)
    document.add_paragraph(nombre)
    document.add_paragraph(f'Fecha: {today.day} de {today.month} de {today.year}')

    for parrafo in parrafos:
        document.add_heading(parrafo, level=1)
        document.add_paragraph(chat_gpt(parrafo))

    # Estilos de fuente
    font_styles = document.styles
    font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    font_object = font_charstyle.font
    font_object.size = Pt(12)
    font_object.name = 'Times New Roman'

    # Guardar documento
    document.save(f'{titulo}.docx')

def create_word_document(parrafos, documento):
    try:

        # Crea un nuevo archivo de Word
        doc = docx.Document()

        for parrafo in parrafos:
            # Agrega un párrafo al archivo
            doc.add_paragraph(f"{parrafo}")
            doc.add_paragraph(chat_gpt(parrafo))

            # Guarda el archivo de Word
        doc.save(f"{documento}.docx")
    except Exception as e:
        print(e)

